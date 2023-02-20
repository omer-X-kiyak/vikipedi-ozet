import wikipedia
import docx
import os
import time

# Wikipedia dil ayarını Türkçe olarak ayarla
wikipedia.set_lang("tr")

while True:
    # Logo Text-to-ASCII-Art-Generator (TAGG)
    print(r"""
       /$ /$   /$$      /$$ /$$$$$$$$ /$$$$$$$        /$$   /$$ /$$$$$$ /$$     /$$ /$$$$$$  /$$   /$$
      |_/|_/  | $$$    /$$$| $$_____/| $$__  $$      | $$  /$$/|_  $$_/|  $$   /$$//$$__  $$| $$  /$$/
      /$$$$$$ | $$$$  /$$$$| $$      | $$  \ $$      | $$ /$$/   | $$   \  $$ /$$/| $$  \ $$| $$ /$$/
     /$$__  $$| $$ $$/$$ $$| $$$$$   | $$$$$$$/      | $$$$$/    | $$    \  $$$$/ | $$$$$$$$| $$$$$/
    | $$  \ $$| $$  $$$| $$| $$__/   | $$__  $$      | $$  $$    | $$     \  $$/  | $$__  $$| $$  $$
    | $$  | $$| $$\  $ | $$| $$      | $$  \ $$      | $$\  $$   | $$      | $$   | $$  | $$| $$\  $$
    |  $$$$$$/| $$ \/  | $$| $$$$$$$$| $$  | $$      | $$ \  $$ /$$$$$$    | $$   | $$  | $$| $$ \ $$
     \______/ |__/     |__/|________/|__/  |__/      |__/  \__/|______/    |__/   |__/  |__/|__/  \__/
    """)
    print("\n***************************************************************** *")
    print("\n* Copyright of Ömer KIYAK, 2022                                 * *")
    print("\n* https://www.instagram.com/omer_x_kiyak/                       * *")
    print("\n* https://github.com/omer-X-kiyak/                              * *")
    print("\n***************************************************************** *\n \n")

    # Arama terimini kullanıcıdan al
    check = input("Ne hakkında bilgi almak istersiniz (Çıkmak için q): ")

    if check.lower() == "q":
        break

    result = wikipedia.search(check)

    # Yeni bir Word belgesi oluştur
    doc = docx.Document()

    # core_properties özelliğine ait author alanını ayarla
    doc.core_properties.author = "omer-x-kiyak"

    # Her bir arama sonucu için özetleri belgeye ekle
    for item in result:
        try:
            # Sayfayı getir ve özetini ekle
            page = wikipedia.page(item)
            if page.title:
                doc.add_heading(page.title, level=1)
                doc.add_paragraph(page.summary)
        except wikipedia.exceptions.DisambiguationError as e:
            print(e.options)
            continue
        except wikipedia.exceptions.PageError as e:
            print(f"Sayfa bulunamadı: {e}")
            continue

    clear = lambda: os.system('cls')
    clear()
  
   
    if result:
        # Dosyayı kaydet
        doc.save(check + ' kısa_arama_sonucu.docx')
        time.sleep(4)
        print("Dosya kaydedildi.")
    else:
        print("Arama sonucu bulunamadı.")
